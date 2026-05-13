#!/usr/bin/env python3
"""Генерирует docs/KurashNation-Moduli-i-Roli.docx (нужен пакет python-docx из venv)."""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
except ImportError:
    print("Установите: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "KurashNation-Moduli-i-Roli.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    t = doc.add_paragraph()
    r = t.add_run("Kurash Nation — описание модулей и ролей")
    r.bold = True
    r.font.size = Pt(16)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_para(
        doc,
        "Документ описывает веб-приложение (Vue 3 + Spring Boot + PostgreSQL): назначение модулей, "
        "возможности по ролям и общую логику работы. REST API доступен по префиксу /api; "
        "аутентификация — JWT (Bearer).",
    )

    add_heading(doc, "1. Архитектура", level=1)
    add_bullets(
        doc,
        [
            "Frontend: Vue 3 (Composition API), маршрутизация, Pinia, Axios к http://host:8080/api (переменная VITE_API_BASE_URL).",
            "Backend: Spring Boot 3, JPA, Spring Security, Liquibase (миграции БД), OpenAPI/Swagger.",
            "Данные: PostgreSQL (jsonb для планов и журналов).",
            "Файлы: загрузки (достижения, справочник) на диске в каталоге storage (см. application.yml).",
        ],
    )

    add_heading(doc, "2. Роли", level=1)
    add_para(doc, "ATHLETE (спортсмен)", bold=True)
    add_bullets(
        doc,
        [
            "Личный кабинет: дашборд, рекомендации, журнал тренировок, журнал достижений, галерея медиа по достижениям.",
            "Видит план недели и совет (tip), в том числе назначенный тренером из справочника (custom_tip).",
        ],
    )
    add_para(doc, "COACH (тренер)", bold=True)
    add_bullets(
        doc,
        [
            "Кабинет и рабочая область: обзор закреплённых спортсменов, инсайты по достижениям (в т.ч. ИИ при настроенном API-ключе).",
            "Справочник: чтение документов администратора, скачивание с JWT, предпросмотр (PDF — в браузере; DOC/DOCX/XLS — извлечённый текст).",
            "ИИ по документу справочника: генерация недельного плана и назначение его выбранному своему спортсмену (сохраняется как рекомендация).",
            "Работа с тренировками и заявками спортсменов в рамках своей группы.",
        ],
    )
    add_para(doc, "ADMIN (администратор)", bold=True)
    add_bullets(
        doc,
        [
            "Пользователи: просмотр, активация, назначение тренера спортсмену.",
            "Справочник: загрузка/редактирование метаданных/удаление файлов (PDF, Word, Excel).",
            "ИИ: просмотр/редактирование текстов промптов (талица ai_prompt_settings), статус интеграции (ключ API только на сервере).",
        ],
    )

    add_heading(doc, "3. Маршруты фронтенда (основные)", level=1)
    add_para(doc, "Спортсмен:", bold=True)
    add_bullets(
        doc,
        [
            "/dashboard — дашборд",
            "/recommendations — рекомендации и план недели",
            "/my-trainings — мои тренировки",
            "/achievement-journal — журнал достижений и недельных расписаний (вход для ИИ-инсайтов)",
            "/my-achievements-media — фото/видео к достижениям",
            "/profile — профиль",
        ],
    )
    add_para(doc, "Тренер:", bold=True)
    add_bullets(
        doc,
        [
            "/coach/cabinet — кабинет (карточки спортсменов)",
            "/coach/workspace — рабочая зона (инсайты, когорта)",
            "/coach/library — справочник",
            "/profile — профиль",
        ],
    )
    add_para(doc, "Администратор:", bold=True)
    add_bullets(
        doc,
        [
            "/admin/users — пользователи и роли",
            "/admin/library — справочник (загрузка файлов)",
            "/admin/ai — промпты ИИ и справка по ключу API",
            "/profile — профиль",
        ],
    )

    add_heading(doc, "4. Backend: контроллеры и назначение", level=1)
    rows = [
        ("AuthController (/auth)", "Вход, регистрация спортсмена/тренера, /me, обновление профиля."),
        ("AthleteController (/athletes)", "Профиль спортсмена, список спортсменов тренера."),
        ("CoachController (/coaches)", "Профиль тренера (/me)."),
        ("TrainingController (/trainings)", "CRUD тренировок (спортсмен — свои; тренер — свои атлеты)."),
        ("TrainingRequestSubmissionController", "Заявки спортсмена на добавление тренировки и очередь на проверку тренеру."),
        ("RecommendationController (/recommendations)", "Последняя рекомендация, генерация, план недели, сравнение с эталоном, PDF; тренер генерирует для атлета."),
        ("AchievementController (/achievements)", "Журнал достижений, недели, запрос ИИ-инсайта по когорте тренера."),
        ("AchievementMediaController", "Загрузка и выдача медиафайлов по достижениям."),
        ("TrainingLibraryController (/training-library)", "Список файлов, скачивание/inline, preview-text, analyze, apply-handbook."),
        ("TrainingLibraryAdminController (/admin/training-library)", "Админ: загрузка, обновление меты, удаление."),
        ("AdminController (/admin)", "Пользователи, коучи, назначение тренера атлету."),
        ("AdminAiPromptController (/admin/ai/prompts)", "Список/обновление промптов."),
        ("AdminAiIntegrationController", "Информация о конфигурации ИИ (без секретов)."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    add_heading(doc, "5. Ключевые сервисы (логика)", level=1)
    add_bullets(
        doc,
        [
            "AuthService — аутентификация, JWT, регистрация, привязка профилей athlete/coach.",
            "AthleteService / CoachService — данные профилей и списки.",
            "TrainingService — тренировки и проверка, что тренер управляет только своими спортсменами.",
            "RecommendationService — расчёт дефицитов к gold standard, недельный план, сохранение рекомендаций; saveHandbookPlanForAthlete — план из справочника + custom_tip.",
            "TrainingLibraryService — хранение файлов справочника; TrainingLibraryHandbookService — извлечение текста, вызов OpenAI Responses API, применение плана.",
            "AchievementService / AchievementInsightService — журнал достижений и LLM-инсайты (промпт из БД).",
            "AchievementMediaService — файлы галереи.",
            "AiPromptSettingService — чтение/запись текстов промптов.",
            "UserService — админ-операции над пользователями.",
            "PdfGeneratorService — экспорт рекомендации в PDF.",
            "OpenAiResponsesClient — HTTP-клиент к OpenAI /v1/responses.",
        ],
    )

    add_heading(doc, "6. Справочник и ИИ (поток)", level=1)
    add_bullets(
        doc,
        [
            "Админ загружает файл в training_library_files (типы: PDF, DOC, DOCX, XLS, XLSX).",
            "Тренер получает список и может: скачать (blob + Authorization), открыть PDF в модалке или текст DOC/XLS через preview-text.",
            "analyze: текст документа (до лимита) + опционально контекст спортсмена и комментарий тренера → промпт HANDBOOK_WEEK_PLAN_JSON → JSON summary + weekPlan.",
            "apply-handbook: сохраняется новая запись recommendations с week_plan, пустым массивом deficits, custom_tip (текст для спортсмена) и ссылкой на документ.",
            "Спортсмен в /recommendations видит актуальную рекомендацию (последняя по дате) с этим советом и планом по дням.",
        ],
    )

    add_heading(doc, "7. Рекомендации и эталон (gold standard)", level=1)
    add_bullets(
        doc,
        [
            "Сравнение метрик спортсмена (тренировки, интенсивность, подтягивания и т.д.) с эталоном по весовой категории.",
            "Автоматический недельный план и совет строятся эвристикой по дефицитам, если нет custom_tip.",
            "План из справочника задаётся тренером отдельно и не пересчитывает эталон автоматически (progressPercent при сохранении условный, например 50%).",
        ],
    )

    add_heading(doc, "8. Разработка и DevTools", level=1)
    add_para(
        doc,
        "При использовании spring-boot-devtools для извлечения DOCX настроен файл "
        "META-INF/spring-devtools.properties (restart.include для JAR Apache POI), иначе возможен ClassNotFoundException для XWPFDocument до полного перезапуска.",
    )

    add_heading(doc, "9. Итог", level=1)
    add_para(
        doc,
        "Система связывает три роли: спортсмен вводит данные и получает планы; тренер анализирует когорту, "
        "назначает планы и использует справочник с ИИ; администратор управляет пользователями, файлами и текстами промптов. "
        "Все защищённые запросы требуют JWT, кроме логина и регистрации.",
    )

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
